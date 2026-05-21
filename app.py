import os
import streamlit as st
import sqlite3
import pandas as pd
import joblib
import io
import textwrap
from utils import check_interaction
from datetime import datetime, date
import hashlib
import database

BASE_DIR = os.path.dirname(__file__)
DB_PATH = os.path.join(BASE_DIR, 'healthcare.db')
MODEL_PATH = BASE_DIR
DATA_PATH = BASE_DIR

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ─── Page Config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="AI Based Personalized Healthcare Advisor",
    page_icon="🏥",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─── Database Init ──────────────────────────────────────────────────────────────
def get_db():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    cursor = conn.cursor()
    cursor.executescript("""
    CREATE TABLE IF NOT EXISTS users(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT NOT NULL,
        email TEXT UNIQUE NOT NULL,
        password TEXT NOT NULL,
        age INTEGER,
        gender TEXT,
        role TEXT DEFAULT 'patient',
        phone TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP
    );
    CREATE TABLE IF NOT EXISTS doctors(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER UNIQUE,
        specialization TEXT,
        license_no TEXT UNIQUE,
        hospital TEXT,
        experience_years INTEGER,
        available_days TEXT,
        consultation_fee REAL DEFAULT 500,
        rating REAL DEFAULT 0.0,
        total_reviews INTEGER DEFAULT 0,
        FOREIGN KEY(user_id) REFERENCES users(id)
    );
    CREATE TABLE IF NOT EXISTS patients(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        blood_pressure TEXT,
        blood_sugar INTEGER,
        allergies TEXT,
        past_diseases TEXT,
        current_symptoms TEXT,
        previous_drugs TEXT,
        diagnosed_disease TEXT,
        recommended_drug TEXT,
        alternative_drug TEXT,
        side_effects TEXT,
        risk_level TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(user_id) REFERENCES users(id)
    );
    CREATE TABLE IF NOT EXISTS appointments(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_id INTEGER,
        doctor_id INTEGER,
        appointment_date TEXT,
        appointment_time TEXT,
        reason TEXT,
        status TEXT DEFAULT 'pending',
        notes TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(patient_id) REFERENCES users(id),
        FOREIGN KEY(doctor_id) REFERENCES users(id)
    );
    CREATE TABLE IF NOT EXISTS prescriptions(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        patient_id INTEGER,
        doctor_id INTEGER,
        drug TEXT,
        dosage TEXT,
        frequency TEXT,
        duration TEXT,
        notes TEXT,
        date_prescribed TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(patient_id) REFERENCES patients(id)
    );
    CREATE TABLE IF NOT EXISTS doctor_notes(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        appointment_id INTEGER,
        doctor_id INTEGER,
        patient_user_id INTEGER,
        diagnosis TEXT,
        treatment_plan TEXT,
        follow_up_date TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(appointment_id) REFERENCES appointments(id)
    );
    CREATE TABLE IF NOT EXISTS reviews(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        doctor_id INTEGER,
        patient_id INTEGER,
        rating INTEGER,
        comment TEXT,
        created_at TEXT DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY(doctor_id) REFERENCES doctors(id),
        FOREIGN KEY(patient_id) REFERENCES users(id)
    );
    """)
    conn.commit()
    conn.close()


def ensure_column_exists(conn, table_name, column_def):
    cursor = conn.cursor()
    cursor.execute(f"PRAGMA table_info({table_name})")
    existing_cols = [row[1] for row in cursor.fetchall()]
    column_name = column_def.split()[0]
    if column_name not in existing_cols:
        cursor.execute(f"ALTER TABLE {table_name} ADD COLUMN {column_def}")
        conn.commit()


init_db()

conn = get_db()
ensure_column_exists(conn, "prescriptions", "doctor_id INTEGER")
ensure_column_exists(conn, "prescriptions", "frequency TEXT")
ensure_column_exists(conn, "prescriptions", "duration TEXT")
ensure_column_exists(conn, "prescriptions", "notes TEXT")
ensure_column_exists(conn, "prescriptions", "date_prescribed TEXT DEFAULT CURRENT_TIMESTAMP")
ensure_column_exists(conn, "patients", "diagnosed_disease TEXT")
ensure_column_exists(conn, "patients", "recommended_drug TEXT")
ensure_column_exists(conn, "patients", "alternative_drug TEXT")
ensure_column_exists(conn, "patients", "side_effects TEXT")
ensure_column_exists(conn, "patients", "risk_level TEXT")
conn.close()

# ─── Helpers ───────────────────────────────────────────────────────────────────
def hash_password(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

def safe_transform(encoder, value):
    if value is None:
        value = 'None'
    value = str(value).strip() or 'None'
    return encoder.transform([value])[0] if value in encoder.classes_ else 0


def normalize_drug_name(drug_name):
    if not drug_name:
        return ''
    text = str(drug_name).lower().strip()
    for token in ['tablet', 'capsule', 'mg', 'ml', 'syrup', 'drop', 'drops', 'oral', 'injection', 'tab', 'cap']:
        text = text.replace(token, '')
    text = ''.join(ch for ch in text if ch.isalnum() or ch.isspace()).strip()
    return ' '.join(text.split())


def disease_to_specialty(disease):
    mapping = {
        'Asthma': ['Pulmonologist', 'Respiratory', 'General Physician'],
        'Cold': ['General Physician', 'ENT', 'Internal Medicine'],
        'Diabetes': ['Endocrinologist', 'General Physician'],
        'Fever': ['General Physician', 'Internal Medicine'],
        'Heart Disease': ['Cardiologist'],
        'Hypertension': ['Cardiologist', 'Nephrologist', 'General Physician'],
        'Kidney Disease': ['Nephrologist', 'Urologist', 'General Physician'],
        'Migraine': ['Neurologist', 'General Physician'],
    }
    return mapping.get(disease, ['General Physician'])


def clean_doctor_display_name(name):
    if not name:
        return ''
    cleaned = name.strip()
    if cleaned.lower().startswith('dr. '):
        cleaned = cleaned[4:].strip()
    elif cleaned.lower().startswith('dr '):
        cleaned = cleaned[3:].strip()
    return cleaned or name


def get_doctors_for_specialties(conn, specialties):
    if not specialties:
        return []
    clauses = []
    params = []
    for spec in specialties:
        spec_text = spec.lower()
        clauses.append("LOWER(d.specialization) LIKE ?")
        params.append(f"%{spec_text}%")
    query = f"""
        SELECT u.name, d.specialization, d.hospital, d.experience_years,
               d.available_days, d.consultation_fee, d.rating
        FROM doctors d JOIN users u ON d.user_id=u.id
        WHERE {' OR '.join(clauses)}
        ORDER BY d.rating DESC, d.experience_years DESC
    """
    return conn.execute(query, params).fetchall()


def get_latest_patient_record_id(user_id):
    conn = get_db()
    record = conn.execute("SELECT id FROM patients WHERE user_id=? ORDER BY id DESC LIMIT 1", (user_id,)).fetchone()
    conn.close()
    return record['id'] if record else None


def get_latest_patient_record(user_id):
    conn = get_db()
    record = conn.execute("SELECT * FROM patients WHERE user_id=? ORDER BY id DESC LIMIT 1", (user_id,)).fetchone()
    conn.close()
    return dict(record) if record else None


def generate_patient_report(patient, prescriptions):
    enriched_patient = dict(patient)
    if patient.get('id') is not None:
        conn = get_db()
        db_patient = conn.execute("SELECT * FROM patients WHERE id=?", (patient['id'],)).fetchone()
        if db_patient:
            for key, value in dict(db_patient).items():
                if enriched_patient.get(key) in (None, '', 'Not available'):
                    enriched_patient[key] = value
        if enriched_patient.get('user_id') is not None:
            user_row = conn.execute("SELECT name, age, gender, phone FROM users WHERE id=?", (enriched_patient['user_id'],)).fetchone()
            if user_row:
                for key, value in dict(user_row).items():
                    if enriched_patient.get(key) in (None, '', 'Not available'):
                        enriched_patient[key] = value
        conn.close()
    elif enriched_patient.get('user_id') is not None:
        conn = get_db()
        user_row = conn.execute("SELECT name, age, gender, phone FROM users WHERE id=?", (enriched_patient['user_id'],)).fetchone()
        if user_row:
            for key, value in dict(user_row).items():
                if enriched_patient.get(key) in (None, '', 'Not available'):
                    enriched_patient[key] = value
        conn.close()

    patient = enriched_patient
    patient_name = patient.get('name') or 'Patient'
    diagnosis = patient.get('diagnosed_disease') or 'Not available'
    recommended = patient.get('recommended_drug') or 'Not available'
    alternative = patient.get('alternative_drug') or 'Not available'
    side_effects = patient.get('side_effects') or 'Not available'
    risk_level = patient.get('risk_level') or 'Not available'
    created_at = str(patient.get('record_date') or patient.get('created_at') or 'Not available')[:10]
    report_title = f"Patient Health Report - {patient_name}"
    report_date = datetime.now().strftime('%Y-%m-%d %H:%M')

    if PDF_AVAILABLE:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font('Helvetica', 'B', 16)
        pdf.cell(0, 10, report_title, ln=True)
        pdf.ln(4)
        pdf.set_font('Helvetica', '', 11)
        pdf.multi_cell(0, 8, f"Generated on: {report_date}")
        pdf.ln(2)
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, 'Patient Information', ln=True)
        pdf.set_font('Helvetica', '', 11)
        pdf.multi_cell(0, 7, f"Name: {patient_name}\nAge: {patient['age']} yrs\nGender: {patient['gender']}\nPhone: {patient['phone'] or 'N/A'}\nLatest record date: {created_at}")
        pdf.ln(2)
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, 'Clinical Summary', ln=True)
        pdf.set_font('Helvetica', '', 11)
        pdf.multi_cell(0, 7, f"Blood Pressure: {patient['blood_pressure'] or 'N/A'}\nBlood Sugar: {patient['blood_sugar'] or 'N/A'} mg/dL\nAI Diagnosis: {diagnosis}\nRisk Level: {risk_level}\nRecommended Drug: {recommended}\nAlternative Drug: {alternative}\nSide Effects: {side_effects}")
        pdf.ln(2)
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, 'Doctor Prescriptions', ln=True)
        pdf.set_font('Helvetica', '', 11)
        if prescriptions:
            for idx, row in enumerate(prescriptions, start=1):
                notes_line = f"\n   Notes: {row['notes']}" if row['notes'] else ''
                pdf.multi_cell(0, 7, f"{idx}. {row['drug']}\n   Dosage: {row['dosage'] or 'N/A'}\n   Frequency: {row['frequency'] or 'N/A'}\n   Duration: {row['duration'] or 'N/A'}\n   Prescribed on: {row['date_prescribed'][:10] if row['date_prescribed'] else 'N/A'}{notes_line}")
                pdf.ln(1)
        else:
            pdf.multi_cell(0, 7, 'No doctor prescriptions found.')

        pdf_bytes = pdf.output(dest='S').encode('latin-1')
        return pdf_bytes, f"{patient_name.replace(' ', '_')}_health_report.pdf", 'application/pdf'

    if DOCX_AVAILABLE:
        doc = Document()
        doc.add_heading(report_title, level=1)
        doc.add_paragraph(f"Generated on: {report_date}")
        doc.add_paragraph(f"Name: {patient_name}")
        doc.add_paragraph(f"Age: {patient['age']} yrs")
        doc.add_paragraph(f"Gender: {patient['gender']}")
        doc.add_paragraph(f"Phone: {patient['phone'] or 'N/A'}")
        doc.add_paragraph(f"Latest record date: {created_at}")
        doc.add_heading('Clinical Summary', level=2)
        doc.add_paragraph(f"Blood Pressure: {patient['blood_pressure'] or 'N/A'}")
        doc.add_paragraph(f"Blood Sugar: {patient['blood_sugar'] or 'N/A'} mg/dL")
        doc.add_paragraph(f"AI Diagnosis: {diagnosis}")
        doc.add_paragraph(f"Risk Level: {risk_level}")
        doc.add_paragraph(f"Recommended Drug: {recommended}")
        doc.add_paragraph(f"Alternative Drug: {alternative}")
        doc.add_paragraph(f"Side Effects: {side_effects}")
        doc.add_heading('Doctor Prescriptions', level=2)
        if prescriptions:
            for idx, row in enumerate(prescriptions, start=1):
                doc.add_paragraph(f"{idx}. {row['drug']}", style='List Number')
                doc.add_paragraph(f"   Dosage: {row['dosage'] or 'N/A'}")
                doc.add_paragraph(f"   Frequency: {row['frequency'] or 'N/A'}")
                doc.add_paragraph(f"   Duration: {row['duration'] or 'N/A'}")
                doc.add_paragraph(f"   Prescribed on: {row['date_prescribed'][:10] if row['date_prescribed'] else 'N/A'}")
                if row['notes']:
                    doc.add_paragraph(f"   Notes: {row['notes']}")
        else:
            doc.add_paragraph('No doctor prescriptions found.')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue(), f"{patient_name.replace(' ', '_')}_health_report.docx", 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

    lines = [
        report_title,
        f"Generated on: {report_date}",
        "",
        "Patient Information",
        f"Name: {patient_name}",
        f"Age: {patient['age']} yrs",
        f"Gender: {patient['gender']}",
        f"Phone: {patient['phone'] or 'N/A'}",
        f"Latest record date: {created_at}",
        "",
        "Clinical Summary",
        f"Blood Pressure: {patient['blood_pressure'] or 'N/A'}",
        f"Blood Sugar: {patient['blood_sugar'] or 'N/A'} mg/dL",
        f"AI Diagnosis: {diagnosis}",
        f"Risk Level: {risk_level}",
        f"Recommended Drug: {recommended}",
        f"Alternative Drug: {alternative}",
        f"Side Effects: {side_effects}",
        "",
        "Doctor Prescriptions",
    ]
    if prescriptions:
        for idx, row in enumerate(prescriptions, start=1):
            lines.extend([
                f"{idx}. {row['drug']}",
                f"   Dosage: {row['dosage'] or 'N/A'}",
                f"   Frequency: {row['frequency'] or 'N/A'}",
                f"   Duration: {row['duration'] or 'N/A'}",
                f"   Prescribed on: {row['date_prescribed'][:10] if row['date_prescribed'] else 'N/A'}",
            ])
            if row['notes']:
                lines.append(f"   Notes: {row['notes']}")
            lines.append("")
    else:
        lines.append('No doctor prescriptions found.')

    text_bytes = '\n'.join(lines).encode('utf-8')
    return text_bytes, f"{patient_name.replace(' ', '_')}_health_report.txt", 'text/plain'


def build_feature_vector(record, user_info):
    bp = str(record.get('blood_pressure', '') or '')
    sys_, dia_ = 0, 0
    if '/' in bp:
        parts = [p.strip() for p in bp.split('/')]
        if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
            sys_, dia_ = int(parts[0]), int(parts[1])
    return [
        int(user_info.get('age', 0) or 0),
        safe_transform(encoders['gender'], user_info.get('gender')),
        sys_,
        dia_,
        int(record.get('blood_sugar', 0) or 0),
        safe_transform(encoders['allergies'], record.get('allergies')),
        safe_transform(encoders['past_diseases'], record.get('past_diseases')),
        safe_transform(encoders['current_symptoms'], record.get('current_symptoms')),
        safe_transform(encoders['previous_drugs'], record.get('previous_drugs')),
    ]


def recommend_drugs_for_disease(disease):
    rows = healthcare_df[healthcare_df['diagnosed_disease'] == disease]
    if rows.empty:
        return None, None, None
    primary = rows['recommended_drug'].mode()[0]
    alternate = rows['alternative_drug'].mode()[0]
    side_fx = rows['side_effects'].mode()[0]
    return primary, alternate, side_fx


def compare_ai_vs_doctor_prescriptions(conn, patient_id, ai_primary, ai_alternative):
    rows = conn.execute("""
        SELECT drug, dosage, frequency, duration, notes, date_prescribed
        FROM prescriptions
        WHERE patient_id=? AND doctor_id IS NOT NULL
        ORDER BY date_prescribed DESC
    """, (patient_id,)).fetchall()
    comparison = []
    for row in rows:
        status = 'Match' if str(row['drug']).strip().lower() in [str(ai_primary).strip().lower(), str(ai_alternative).strip().lower()] else 'Different'
        comparison.append({**dict(row), 'match_status': status})
    return comparison


def calculate_dataset_model_accuracy():
    if model is None or healthcare_df.empty:
        return None
    total = 0
    correct = 0
    for _, row in healthcare_df.iterrows():
        try:
            bp = str(row['blood_pressure'])
            if '/' not in bp:
                continue
            parts = [p.strip() for p in bp.split('/')]
            if len(parts) != 2 or not parts[0].isdigit() or not parts[1].isdigit():
                continue
            sys_, dia_ = int(parts[0]), int(parts[1])
            feats = [
                int(row.get('age', 0) or 0),
                safe_transform(encoders['gender'], row.get('gender')),
                sys_, dia_,
                int(row.get('blood_sugar', 0) or 0),
                safe_transform(encoders['allergies'], row.get('allergies')),
                safe_transform(encoders['past_diseases'], row.get('past_diseases')),
                safe_transform(encoders['current_symptoms'], row.get('current_symptoms')),
                safe_transform(encoders['previous_drugs'], row.get('previous_drugs')),
            ]
            pred = model.predict([feats])[0]
            predicted_disease = encoders['diagnosed_disease'].inverse_transform([pred])[0]
            if predicted_disease == row['diagnosed_disease']:
                correct += 1
            total += 1
        except Exception:
            continue
    return (correct, total, correct / total) if total > 0 else None

# ─── Load ML Model ─────────────────────────────────────────────────────────────
@st.cache_resource
def load_model():
    try:
        model = joblib.load(os.path.join(MODEL_PATH, 'disease_model.pkl'))
        encoders = {
            'gender': joblib.load(os.path.join(MODEL_PATH, 'le_gender.pkl')),
            'allergies': joblib.load(os.path.join(MODEL_PATH, 'le_allergies.pkl')),
            'past_diseases': joblib.load(os.path.join(MODEL_PATH, 'le_past_diseases.pkl')),
            'current_symptoms': joblib.load(os.path.join(MODEL_PATH, 'le_current_symptoms.pkl')),
            'previous_drugs': joblib.load(os.path.join(MODEL_PATH, 'le_previous_drugs.pkl')),
            'diagnosed_disease': joblib.load(os.path.join(MODEL_PATH, 'le_diagnosed_disease.pkl')),
        }
        return model, encoders
    except Exception as e:
        return None, {}

@st.cache_data
def load_healthcare_df():
    try:
        return pd.read_csv(os.path.join(DATA_PATH, 'healthcare_main_dataset_900.csv'))
    except:
        return pd.DataFrame()

model, encoders = load_model()
healthcare_df = load_healthcare_df()

# ─── Session State ─────────────────────────────────────────────────────────────
for key, val in [('logged_in', False), ('user_id', None), ('user_name', ''), ('user_role', 'patient')]:
    if key not in st.session_state:
        st.session_state[key] = val

# ─── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<link href="https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:ital,wght@0,300;0,400;0,500;1,400&display=swap" rel="stylesheet">
<style>
:root {
    --bg: #0a0f1e;
    --surface: #111827;
    --surface2: #1a2234;
    --border: #1e2d45;
    --accent: #00d4ff;
    --accent2: #7c3aed;
    --accent3: #10b981;
    --danger: #ef4444;
    --warning: #f59e0b;
    --text: #e2e8f0;
    --muted: #64748b;
    --card-glow: 0 0 40px rgba(0,212,255,0.06);
}
html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    background: var(--bg) !important;
    color: var(--text) !important;
}
.stApp { background: var(--bg) !important; }
section[data-testid="stSidebar"] {
    background: var(--surface) !important;
    border-right: 1px solid var(--border) !important;
}
.mc-card {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 16px;
    padding: 24px 28px;
    margin-bottom: 20px;
    box-shadow: var(--card-glow);
}
.mc-card:hover { border-color: rgba(0,212,255,0.2); }
.hero-wrap {
    background: linear-gradient(135deg, #0a0f1e 0%, #0f1a2e 50%, #0a0f1e 100%);
    border: 1px solid var(--border);
    border-radius: 20px;
    padding: 40px 48px;
    margin-bottom: 28px;
    position: relative;
    overflow: hidden;
}
.hero-eyebrow {
    font-family: 'DM Sans', sans-serif;
    font-size: 0.72rem;
    font-weight: 500;
    letter-spacing: 0.18em;
    text-transform: uppercase;
    color: var(--accent);
    margin-bottom: 12px;
}
.hero-title {
    font-family: 'Syne', sans-serif;
    font-size: 2.6rem;
    font-weight: 800;
    line-height: 1.15;
    color: #f1f5f9;
    margin-bottom: 14px;
}
.hero-title span { color: var(--accent); }
.hero-sub {
    font-size: 1rem;
    color: var(--muted);
    line-height: 1.7;
    max-width: 600px;
}
.sec-title {
    font-family: 'Syne', sans-serif;
    font-size: 1.4rem;
    font-weight: 700;
    color: #f1f5f9;
    margin-bottom: 6px;
}
.sec-sub {
    font-size: 0.9rem;
    color: var(--muted);
    margin-bottom: 20px;
}
.stat-card {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: 14px;
    padding: 20px 24px;
    text-align: center;
}
.stat-num {
    font-family: 'Syne', sans-serif;
    font-size: 2.2rem;
    font-weight: 800;
    color: var(--accent);
}
.stat-label {
    font-size: 0.82rem;
    color: var(--muted);
    text-transform: uppercase;
    letter-spacing: 0.1em;
    margin-top: 4px;
}
.badge {
    display: inline-block;
    padding: 3px 10px;
    border-radius: 20px;
    font-size: 0.75rem;
    font-weight: 600;
    letter-spacing: 0.04em;
}
.badge-green  { background: rgba(16,185,129,0.15); color: #10b981; border: 1px solid rgba(16,185,129,0.3); }
.badge-red    { background: rgba(239,68,68,0.15);  color: #ef4444; border: 1px solid rgba(239,68,68,0.3); }
.badge-yellow { background: rgba(245,158,11,0.15); color: #f59e0b; border: 1px solid rgba(245,158,11,0.3); }
.badge-blue   { background: rgba(0,212,255,0.12);  color: #00d4ff; border: 1px solid rgba(0,212,255,0.25); }
.badge-purple { background: rgba(124,58,237,0.15); color: #a78bfa; border: 1px solid rgba(124,58,237,0.3); }
.appt-row {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: 12px;
    padding: 16px 20px;
    margin-bottom: 10px;
    display: flex;
    align-items: center;
    gap: 16px;
}
.doc-card {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 16px;
    padding: 22px;
    margin-bottom: 14px;
}
.doc-card:hover {
    border-color: rgba(0,212,255,0.3);
    box-shadow: 0 8px 32px rgba(0,212,255,0.08);
}
.doc-name {
    font-family: 'Syne', sans-serif;
    font-size: 1.1rem;
    font-weight: 700;
    color: #f1f5f9;
}
.doc-spec {
    font-size: 0.85rem;
    color: var(--accent);
    margin: 3px 0 8px;
}
.timeline-item {
    border-left: 2px solid var(--border);
    padding: 0 0 20px 20px;
    position: relative;
    margin-left: 8px;
}
.timeline-dot {
    width: 10px; height: 10px;
    background: var(--accent);
    border-radius: 50%;
    position: absolute;
    left: -6px; top: 4px;
}
.stButton > button {
    background: linear-gradient(135deg, #00d4ff22, #00d4ff11) !important;
    color: var(--accent) !important;
    border: 1px solid rgba(0,212,255,0.35) !important;
    border-radius: 10px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 500 !important;
    padding: 0.55rem 1.1rem !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #00d4ff33, #00d4ff22) !important;
    border-color: rgba(0,212,255,0.6) !important;
}
.stTextInput > label, .stNumberInput > label,
.stSelectbox > label, .stTextArea > label,
.stDateInput > label, .stTimeInput > label {
    color: #94a3b8 !important;
    font-size: 0.85rem !important;
    font-weight: 500 !important;
    letter-spacing: 0.03em !important;
}
.stTextInput input, .stNumberInput input, .stTextArea textarea {
    background: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    color: var(--text) !important;
}
.stTextInput input:focus, .stNumberInput input:focus, .stTextArea textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 2px rgba(0,212,255,0.12) !important;
}
div[data-baseweb="select"] > div {
    background: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 10px !important;
    color: var(--text) !important;
}
.stSuccess > div { background: rgba(16,185,129,0.1) !important; color: #10b981 !important; border: 1px solid rgba(16,185,129,0.25) !important; border-radius: 10px !important; }
.stError   > div { background: rgba(239,68,68,0.1)  !important; color: #ef4444 !important; border: 1px solid rgba(239,68,68,0.25)  !important; border-radius: 10px !important; }
.stWarning > div { background: rgba(245,158,11,0.1) !important; color: #f59e0b !important; border: 1px solid rgba(245,158,11,0.25) !important; border-radius: 10px !important; }
.stInfo    > div { background: rgba(0,212,255,0.08)  !important; color: var(--accent) !important; border: 1px solid rgba(0,212,255,0.2) !important; border-radius: 10px !important; }
.stTabs [data-baseweb="tab-list"] { background: var(--surface) !important; border-radius: 12px !important; border: 1px solid var(--border) !important; gap: 4px !important; padding: 6px !important; }
.stTabs [data-baseweb="tab"] { background: transparent !important; color: var(--muted) !important; border-radius: 8px !important; font-family: 'DM Sans' !important; }
.stTabs [aria-selected="true"] { background: var(--surface2) !important; color: var(--accent) !important; border: 1px solid var(--border) !important; }
.stDataFrame { border: 1px solid var(--border) !important; border-radius: 12px !important; }
div[data-testid="metric-container"] {
    background: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 14px !important;
    padding: 16px !important;
}
div[data-testid="metric-container"] label { color: var(--muted) !important; font-size: 0.8rem !important; }
div[data-testid="metric-container"] [data-testid="metric-value"] { color: var(--accent) !important; font-family: 'Syne', sans-serif !important; }
div[data-testid="stExpander"] { background: var(--surface2) !important; border: 1px solid var(--border) !important; border-radius: 12px !important; }
.stRadio > div { gap: 8px !important; }
.stRadio label { color: var(--text) !important; }
hr { border-color: var(--border) !important; }
</style>
""", unsafe_allow_html=True)

# ─── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("""
    <div style="padding: 8px 0 20px;">
        <div style="font-family:'Syne',sans-serif;font-size:1.3rem;font-weight:800;color:#f1f5f9;">
            Medi<span style="color:#00d4ff;">Core</span> AI
        </div>
        <div style="font-size:0.72rem;color:#64748b;letter-spacing:0.12em;text-transform:uppercase;">Healthcare Platform</div>
    </div>
    """, unsafe_allow_html=True)

    if st.session_state.logged_in:
        st.markdown(f"""
        <div style="background:#1a2234;border:1px solid #1e2d45;border-radius:12px;padding:14px 16px;margin-bottom:20px;">
            <div style="font-size:0.75rem;color:#64748b;margin-bottom:4px;">Logged in as</div>
            <div style="font-weight:600;color:#f1f5f9;">{st.session_state.user_name}</div>
            <div style="margin-top:6px;">
                <span class="badge {'badge-blue' if st.session_state.user_role == 'doctor' else 'badge-green'}">
                    {'🩺 Doctor' if st.session_state.user_role == 'doctor' else '🧑 Patient'}
                </span>
            </div>
        </div>
        """, unsafe_allow_html=True)

        if st.session_state.user_role == 'patient':
            pages = ["🏠 Dashboard", "📋 Enter Medical Data", "🔬 Analyze Symptoms",
                     "📅 Book Appointment", "👨‍⚕️ Find Doctors", "📁 Medical History", "💊 My Prescriptions"]
        else:
            pages = ["🏠 Dashboard", "📅 My Appointments", "👥 My Patients",
                     "✍️ Write Prescription", "📊 Analytics", "👤 Doctor Profile"]

        page = st.radio("Navigation", pages, label_visibility="collapsed")
        st.markdown("---")
        if st.button("🚪 Sign Out", use_container_width=True):
            for k in ['logged_in', 'user_id', 'user_name', 'user_role']:
                st.session_state[k] = False if k == 'logged_in' else None if 'id' in k else ''
            st.session_state.user_role = 'patient'
            st.rerun()
    else:
        page = "auth"
        st.markdown("""
        <div style="background:#1a2234;border:1px solid #1e2d45;border-radius:12px;padding:16px;font-size:0.88rem;color:#64748b;line-height:1.6;">
            <b style="color:#e2e8f0;">Welcome!</b><br>
            Sign in to access personalized AI-powered healthcare advice, doctor consultations, and medical history tracking.
        </div>
        """, unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# AUTH PAGE
# ══════════════════════════════════════════════════════════════════════════════
if not st.session_state.logged_in:
    st.markdown("""
    <div class="hero-wrap">
        <div class="hero-eyebrow">Next-Gen Clinical AI Platform</div>
        <div class="hero-title">Your health, <span>intelligently</span> managed.</div>
        <div class="hero-sub">MediCore AI connects patients with doctors, predicts conditions with ML, checks drug safety, and tracks your complete medical journey — all in one secure platform.</div>
    </div>
    """, unsafe_allow_html=True)

    feat_col1, feat_col2, feat_col3 = st.columns(3)
    for col, icon, title, desc in [
        (feat_col1, "🔬", "AI Diagnosis", "ML model predicts conditions from symptoms, vitals, and history"),
        (feat_col2, "👨‍⚕️", "Doctor Network", "Book appointments and consult verified specialists"),
        (feat_col3, "💊", "Drug Safety", "Real-time interaction checks before any prescription"),
    ]:
        with col:
            st.markdown(f"""
            <div class="mc-card">
                <div style="font-size:1.8rem;margin-bottom:10px;">{icon}</div>
                <div class="sec-title" style="font-size:1rem;">{title}</div>
                <div style="font-size:0.85rem;color:#64748b;line-height:1.6;">{desc}</div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("---")
    auth_col, _ = st.columns([1.1, 0.9])
    with auth_col:
        tab_login, tab_reg = st.tabs(["  Sign In  ", "  Create Account  "])

        with tab_login:
            st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)
            l_email = st.text_input("Email address", key="li_email", placeholder="you@example.com")
            l_pass  = st.text_input("Password", type="password", key="li_pass", placeholder="••••••••")
            if st.button("Sign In →", use_container_width=True):
                conn = get_db()
                user = conn.execute(
                    "SELECT id, name, role FROM users WHERE email=? AND password=?",
                    (l_email, hash_password(l_pass))
                ).fetchone()
                conn.close()
                if user:
                    st.session_state.logged_in = True
                    st.session_state.user_id   = user['id']
                    st.session_state.user_name = user['name']
                    st.session_state.user_role = user['role']
                    st.rerun()
                else:
                    st.error("Incorrect email or password.")

        with tab_reg:
            st.markdown("<div style='height:16px'></div>", unsafe_allow_html=True)
            role_sel = st.selectbox("Registering as", ["Patient", "Doctor"], key="reg_role")
            rc1, rc2 = st.columns(2)
            with rc1:
                r_name  = st.text_input("Full Name", key="rn")
                r_age   = st.number_input("Age", 1, 120, 30, key="ra")
            with rc2:
                r_email = st.text_input("Email", key="re")
                r_gender= st.selectbox("Gender", ["Male","Female","Other"], key="rg")
            r_phone = st.text_input("Phone", key="rph", placeholder="+91 XXXXX XXXXX")
            r_pass  = st.text_input("Password", type="password", key="rp")

            if role_sel == "Doctor":
                st.markdown("<div style='margin-top:8px;font-size:0.85rem;color:#64748b;'>Doctor details</div>", unsafe_allow_html=True)
                d1, d2 = st.columns(2)
                with d1:
                    spec   = st.text_input("Specialization", key="dspec")
                    hosp   = st.text_input("Hospital / Clinic", key="dhosp")
                    exp    = st.number_input("Years of Experience", 0, 60, 5, key="dexp")
                with d2:
                    lic    = st.text_input("License No.", key="dlic")
                    days   = st.text_input("Available Days (e.g. Mon,Wed,Fri)", key="ddays")
                    fee    = st.number_input("Consultation Fee (₹)", 0, 100000, 500, key="dfee")

            if st.button("Create Account →", use_container_width=True):
                if not r_name or not r_email or not r_pass:
                    st.error("Please fill all required fields.")
                else:
                    conn = get_db()
                    try:
                        role = role_sel.lower()
                        conn.execute(
                            "INSERT INTO users(name,email,password,age,gender,role,phone) VALUES(?,?,?,?,?,?,?)",
                            (r_name, r_email, hash_password(r_pass), r_age, r_gender, role, r_phone)
                        )
                        conn.commit()
                        if role == 'doctor':
                            uid = conn.execute("SELECT id FROM users WHERE email=?", (r_email,)).fetchone()['id']
                            conn.execute(
                                "INSERT INTO doctors(user_id,specialization,license_no,hospital,experience_years,available_days,consultation_fee) VALUES(?,?,?,?,?,?,?)",
                                (uid, spec, lic, hosp, exp, days, fee)
                            )
                            conn.commit()
                        st.success("Account created! Please sign in.")
                    except sqlite3.IntegrityError:
                        st.error("Email or License No. already registered.")
                    conn.close()

# ══════════════════════════════════════════════════════════════════════════════
# PATIENT PAGES
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.user_role == 'patient':

    # ── Dashboard ──────────────────────────────────────────────────────────────
    if "Dashboard" in page:
        conn = get_db()
        rec_count   = conn.execute("SELECT COUNT(*) FROM patients WHERE user_id=?", (st.session_state.user_id,)).fetchone()[0]
        appt_count  = conn.execute("SELECT COUNT(*) FROM appointments WHERE patient_id=?", (st.session_state.user_id,)).fetchone()[0]
        doc_count   = conn.execute("SELECT COUNT(*) FROM doctors").fetchone()[0]
        last_diag   = conn.execute("SELECT diagnosed_disease, risk_level, created_at FROM patients WHERE user_id=? AND diagnosed_disease IS NOT NULL ORDER BY id DESC LIMIT 1", (st.session_state.user_id,)).fetchone()
        next_appt   = conn.execute("SELECT a.appointment_date, a.appointment_time, u.name, a.status FROM appointments a JOIN users u ON a.doctor_id=u.id WHERE a.patient_id=? AND a.appointment_date >= date('now') ORDER BY a.appointment_date LIMIT 1", (st.session_state.user_id,)).fetchone()
        conn.close()

        st.markdown(f"""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Patient Dashboard</div>
            <div class="hero-title">Good day, <span>{st.session_state.user_name.split()[0]}</span> 👋</div>
            <div class="hero-sub">Here's an overview of your health activity on MediCore AI.</div>
        </div>
        """, unsafe_allow_html=True)

        m1, m2, m3, m4 = st.columns(4)
        for col, num, lbl in [(m1, rec_count, "Medical Records"), (m2, appt_count, "Appointments"), (m3, doc_count, "Verified Doctors"), (m4, "Active" if rec_count else "—", "Status")]:
            with col:
                col.metric(lbl, num)

        st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("""<div class="sec-title" style="margin-bottom:12px;">⚡ Quick Actions</div>""", unsafe_allow_html=True)
            qa1, qa2 = st.columns(2)
            with qa1:
                if st.button("📋 Add Medical Data", use_container_width=True): st.info("Go to 'Enter Medical Data' in the sidebar.")
                if st.button("🔬 Run AI Analysis",  use_container_width=True): st.info("Go to 'Analyze Symptoms' in the sidebar.")
            with qa2:
                if st.button("📅 Book Appointment", use_container_width=True): st.info("Go to 'Book Appointment' in the sidebar.")
                if st.button("👨‍⚕️ Find Doctors",     use_container_width=True): st.info("Go to 'Find Doctors' in the sidebar.")

        with c2:
            st.markdown("""<div class="sec-title" style="margin-bottom:12px;">📌 Snapshot</div>""", unsafe_allow_html=True)
            if last_diag:
                risk_badge = {"High":"badge-red","Medium":"badge-yellow","Low":"badge-green"}.get(last_diag['risk_level'], 'badge-blue')
                st.markdown(f"""
                <div class="mc-card mc-card-accent" style="margin-bottom:10px;">
                    <div style="font-size:0.75rem;color:#64748b;margin-bottom:4px;">Last Diagnosis</div>
                    <div style="font-weight:600;color:#f1f5f9;font-size:1.05rem;">{last_diag['diagnosed_disease']}</div>
                    <div style="margin-top:8px;"><span class="badge {risk_badge}">Risk: {last_diag['risk_level']}</span></div>
                </div>
                """, unsafe_allow_html=True)
            if next_appt:
                st.markdown(f"""
                <div class="mc-card mc-card-green">
                    <div style="font-size:0.75rem;color:#64748b;margin-bottom:4px;">Next Appointment</div>
                    <div style="font-weight:600;color:#f1f5f9;">{next_appt['appointment_date']} at {next_appt['appointment_time']}</div>
                    <div style="font-size:0.85rem;color:#64748b;">Dr. {next_appt['name']}</div>
                    <div style="margin-top:8px;"><span class="badge badge-blue">{next_appt['status'].title()}</span></div>
                </div>
                """, unsafe_allow_html=True)
            if not last_diag and not next_appt:
                st.info("No records yet. Start by entering your medical data.")

    # ── Enter Medical Data ─────────────────────────────────────────────────────
    elif "Medical Data" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Data Entry</div>
            <div class="hero-title">Enter <span>Medical</span> Details</div>
            <div class="hero-sub">Provide accurate vitals and history for the best AI-driven analysis results.</div>
        </div>
        """, unsafe_allow_html=True)

        with st.form("patient_form", clear_on_submit=True):
            st.markdown('<div class="sec-title">🫀 Vital Signs</div>', unsafe_allow_html=True)
            v1, v2 = st.columns(2)
            with v1:
                blood_pressure = st.text_input("Blood Pressure", placeholder="120/80")
            with v2:
                blood_sugar = st.number_input("Blood Sugar (mg/dL)", min_value=0, value=90)

            st.markdown('<div class="sec-title" style="margin-top:16px;">📋 Medical History</div>', unsafe_allow_html=True)
            h1, h2 = st.columns(2)
            with h1:
                allergies    = st.text_input("Allergies", placeholder="Penicillin, Dust, Pollen")
                current_syms = st.text_input("Current Symptoms", placeholder="Fever, Cough, Headache")
            with h2:
                past_diseases= st.text_input("Past Diseases", placeholder="Diabetes, Hypertension")
                prev_drugs   = st.text_input("Previous / Current Drugs", placeholder="Metformin, Aspirin")

            submitted = st.form_submit_button("💾 Save Medical Record", use_container_width=True)
            if submitted:
                if '/' not in blood_pressure or not all(p.strip().isdigit() for p in blood_pressure.split('/')):
                    st.error("Enter blood pressure as Systolic/Diastolic (e.g. 120/80).")
                else:
                    conn = get_db()
                    conn.execute("""INSERT INTO patients(user_id,blood_pressure,blood_sugar,allergies,past_diseases,current_symptoms,previous_drugs)
                        VALUES(?,?,?,?,?,?,?)""",
                        (st.session_state.user_id, blood_pressure, blood_sugar, allergies, past_diseases, current_syms, prev_drugs))
                    conn.commit()
                    conn.close()
                    st.success("✅ Medical record saved! Proceed to Analyze Symptoms.")

    # ── Analyze Symptoms ────────────────────────────────────────────────────────
    elif "Analyze" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">AI Engine</div>
            <div class="hero-title"><span>Symptom</span> Analysis</div>
            <div class="hero-sub">Our Random Forest model analyzes your vitals and history to predict conditions, recommend safe medications, and suggest the right specialists.</div>
        </div>
        """, unsafe_allow_html=True)

        patient = get_latest_patient_record(st.session_state.user_id)
        conn = get_db()
        user_info = conn.execute("SELECT name, age, gender, phone FROM users WHERE id=?", (st.session_state.user_id,)).fetchone()
        conn.close()

        if not patient:
            st.warning("⚠️ No medical record found. Please go to 'Enter Medical Data' first.")
        elif model is None:
            st.error("ML model not found. Ensure disease_model.pkl and encoder files are present.")
        else:
            # Show current record
            st.markdown('<div class="sec-title">📄 Current Record</div>', unsafe_allow_html=True)
            rc1, rc2, rc3 = st.columns(3)
            rc1.metric("Blood Pressure", patient.get('blood_pressure') or '—')
            rc2.metric("Blood Sugar", f"{patient.get('blood_sugar') or '—'} mg/dL")
            rc3.metric("Allergies", patient.get('allergies') or 'None')

            st.markdown("<div style='margin-top:12px;'></div>", unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:0.85rem;color:#94a3b8;'>Using complete patient record data from the database: past illnesses, symptoms, allergies, previous drugs, vitals, age, and gender.</div>", unsafe_allow_html=True)
            if st.button("🔬 Run AI Analysis", use_container_width=True):
                if not patient.get('blood_pressure') or '/' not in str(patient.get('blood_pressure')):
                    st.error("Invalid blood pressure in record.")
                else:
                    try:
                        feats = build_feature_vector(patient, dict(user_info))
                        pred = model.predict([feats])[0]
                        disease = encoders['diagnosed_disease'].inverse_transform([pred])[0]

                        st.markdown(f"""
                        <div class="mc-card mc-card-accent" style="margin:20px 0;">
                            <div style="font-size:0.78rem;color:#64748b;margin-bottom:6px;">AI PREDICTION</div>
                            <div style="font-family:'Syne',sans-serif;font-size:1.8rem;font-weight:800;color:#00d4ff;">{disease}</div>
                        </div>
                        """, unsafe_allow_html=True)

                        rows = healthcare_df[healthcare_df['diagnosed_disease'] == disease]
                        if not rows.empty:
                            rec_drug  = rows['recommended_drug'].mode()[0]
                            alt_drug  = rows['alternative_drug'].mode()[0]
                            side_fx   = rows['side_effects'].mode()[0]
                            risk_lvl  = rows['risk_level'].mode()[0]

                            risk_badge = {"High":"badge-red","Medium":"badge-yellow","Low":"badge-green"}.get(risk_lvl, 'badge-blue')

                            ar1, ar2 = st.columns([1.2, 0.8])
                            with ar1:
                                st.markdown(f"""
                                <div class="mc-card mc-card-green">
                                    <div class="sec-title" style="font-size:1rem;margin-bottom:14px;">💊 Drug Recommendations</div>
                                    <div style="margin-bottom:10px;">
                                        <span style="font-size:0.75rem;color:#64748b;">PRIMARY</span><br>
                                        <span style="font-weight:600;color:#f1f5f9;font-size:1.05rem;">{rec_drug}</span>
                                    </div>
                                    <div style="margin-bottom:10px;">
                                        <span style="font-size:0.75rem;color:#64748b;">ALTERNATIVE</span><br>
                                        <span style="font-weight:600;color:#f1f5f9;">{alt_drug}</span>
                                    </div>
                                    <div>
                                        <span style="font-size:0.75rem;color:#64748b;">SIDE EFFECTS</span><br>
                                        <span style="color:#94a3b8;font-size:0.9rem;">{side_fx}</span>
                                    </div>
                                </div>
                                """, unsafe_allow_html=True)
                            with ar2:
                                st.markdown(f"""
                                <div class="mc-card mc-card-purple">
                                    <div class="sec-title" style="font-size:1rem;margin-bottom:14px;">⚠️ Risk Assessment</div>
                                    <div style="margin-bottom:12px;"><span class="badge {risk_badge}" style="font-size:0.9rem;padding:5px 14px;">{risk_lvl} Risk</span></div>
                                    <div style="font-size:0.85rem;color:#64748b;line-height:1.6;">Consult a specialist before starting any medication. This is an AI-generated suggestion only.</div>
                                </div>
                                """, unsafe_allow_html=True)

                            with get_db() as doc_conn:
                                specialist_list = get_doctors_for_specialties(doc_conn, disease_to_specialty(disease))

                            if specialist_list:
                                st.markdown('<div class="sec-title" style="margin:20px 0 12px;">👨‍⚕️ Suggested Specialists Based on Current Diagnosis & History</div>', unsafe_allow_html=True)
                                for doc in specialist_list:
                                    st.markdown(f"""
                                    <div class="doc-card">
                                        <div style="display:flex;justify-content:space-between;align-items:flex-start;gap:12px;">
                                            <div>
                                                <div class="doc-name">Dr. {doc['name']}</div>
                                                <div class="doc-spec">{doc['specialization']}</div>
                                                <div style="font-size:0.85rem;color:#64748b;margin-top:4px;">🏥 {doc['hospital']} · {doc['experience_years']} yrs experience</div>
                                                <div style="font-size:0.85rem;color:#64748b;margin-top:4px;">📅 {doc['available_days'] or 'N/A'}</div>
                                            </div>
                                            <div style="text-align:right;">
                                                <div style="font-family:'Syne',sans-serif;font-size:1.2rem;font-weight:700;color:#00d4ff;">₹{doc['consultation_fee']}</div>
                                                <div style="font-size:0.75rem;color:#64748b;">⭐ {doc['rating'] or 'N/A'}</div>
                                            </div>
                                        </div>
                                    </div>
                                    """, unsafe_allow_html=True)
                            else:
                                st.info("No matching specialists are registered yet. Use the Find Doctors feature to browse available providers.")

                            # Drug interactions
                            st.markdown('<div class="sec-title" style="margin:20px 0 12px;">🔄 Drug Interaction Check</div>', unsafe_allow_html=True)
                            prev = [d.strip() for d in str(patient['previous_drugs'] or '').split(',') if d.strip()]
                            interactions_found = []
                            for d in prev:
                                for target in [rec_drug, alt_drug]:
                                    inter, sev = check_interaction(target, d)
                                    if inter:
                                        interactions_found.append((target, d, inter, sev))
                            if interactions_found:
                                for t, d, inter, sev in interactions_found:
                                    sev_badge = {"High":"badge-red","Moderate":"badge-yellow","Low":"badge-green"}.get(sev, 'badge-blue')
                                    st.markdown(f"""
                                    <div class="mc-card mc-card-red" style="padding:14px 18px;">
                                        <b style="color:#ef4444;">{t}</b> + <b style="color:#f1f5f9;">{d}</b>
                                        <span class="badge {sev_badge}" style="margin-left:10px;">{sev}</span>
                                        <div style="font-size:0.85rem;color:#94a3b8;margin-top:6px;">{inter}</div>
                                    </div>
                                    """, unsafe_allow_html=True)
                            else:
                                st.success("✅ No harmful drug interactions detected.")

                            doctor_prescriptions_conn = get_db()
                            doctor_prescriptions = doctor_prescriptions_conn.execute("""
                                SELECT drug, dosage, frequency, duration, notes, date_prescribed
                                FROM prescriptions
                                WHERE patient_id=? AND doctor_id IS NOT NULL
                                ORDER BY date_prescribed DESC
                            """, (patient['id'],)).fetchall()
                            doctor_prescriptions_conn.close()

                            if doctor_prescriptions:
                                matched = sum(1 for d in doctor_prescriptions if normalize_drug_name(d['drug']) in [normalize_drug_name(rec_drug), normalize_drug_name(alt_drug)])
                                total_doctor = len(doctor_prescriptions)
                                accuracy = (matched / total_doctor) * 100 if total_doctor else 0
                                st.markdown('<div class="sec-title" style="margin:20px 0 12px;">🔍 AI vs Doctor Prescription Review</div>', unsafe_allow_html=True)
                                st.markdown(
                                    "<div style='margin-bottom:12px;color:#94a3b8;'>The AI recommendation is compared against the doctor-issued prescriptions by normalizing drug names and checking whether the prescribed drug matches either the AI primary or alternative suggestion. "
                                    f"Agreement rate: <strong>{matched}/{total_doctor}</strong> ({accuracy:.1f}%).</div>",
                                    unsafe_allow_html=True
                                )
                                for d in doctor_prescriptions[:3]:
                                    match_label = "Match" if str(d['drug']).strip().lower() in [str(rec_drug).strip().lower(), str(alt_drug).strip().lower()] else "Different"
                                    badge = "badge-green" if match_label == "Match" else "badge-yellow"
                                    notes_html = f"<div style='font-size:0.82rem;color:#94a3b8;margin-top:6px;'>📝 {d['notes']}</div>" if d['notes'] else ''
                                    st.markdown(f"""
                                    <div class="mc-card mc-card-accent" style="margin-bottom:10px;">
                                        <div style="display:flex;justify-content:space-between;align-items:flex-start;">
                                            <div>
                                                <div style="font-weight:700;color:#f1f5f9;">{d['drug']}</div>
                                                <div style="font-size:0.85rem;color:#94a3b8;">{d['dosage'] or 'As prescribed'} · {d['frequency'] or '—'} · {d['duration'] or '—'}</div>
                                                {notes_html}
                                            </div>
                                            <div style="text-align:right;">
                                                <span class="badge {badge}">{match_label}</span>
                                                <div style="font-size:0.75rem;color:#64748b;margin-top:6px;">{d['date_prescribed'][:10] if d['date_prescribed'] else '—'}</div>
                                            </div>
                                        </div>
                                    </div>
                                    """, unsafe_allow_html=True)
                            else:
                                st.info("No doctor-issued prescriptions were found for comparison yet.")

                            dataset_accuracy = calculate_dataset_model_accuracy()
                            if dataset_accuracy:
                                corr, total, pct = dataset_accuracy
                                st.markdown('<div class="sec-title" style="margin:20px 0 12px;">📈 Model Training Accuracy</div>', unsafe_allow_html=True)
                                st.markdown(f"<div style='margin-bottom:12px;color:#94a3b8;'>The model correctly predicted the diagnosed disease for <strong>{corr}/{total}</strong> samples from the training dataset ({pct*100:.1f}%).</div>", unsafe_allow_html=True)
                            else:
                                st.info("Model training accuracy is unavailable because the model or training data could not be loaded.")

                            conn = get_db()
                            cur = conn.cursor()
                            cur.execute("""UPDATE patients SET diagnosed_disease=?,recommended_drug=?,alternative_drug=?,side_effects=?,risk_level=? WHERE id=?""",
                                (disease, rec_drug, alt_drug, side_fx, risk_lvl, patient['id']))
                            # verify update
                            conn.commit()
                            saved_row = cur.execute(
                                "SELECT diagnosed_disease,recommended_drug,alternative_drug,side_effects,risk_level,created_at FROM patients WHERE id=?",
                                (patient['id'],)
                            ).fetchone()
                            # Insert AI-suggested prescription record
                            cur.execute("INSERT INTO prescriptions(patient_id,drug,dosage,frequency,duration,date_prescribed) VALUES(?,?,?,?,?,?)",
                                (patient['id'], rec_drug, "As prescribed", "Twice daily", "7 days", datetime.now().strftime("%Y-%m-%d")))
                            conn.commit()
                            # surface confirmation to user and keep in session state for immediacy
                            if saved_row:
                                saved_vals = dict(saved_row)
                                st.success(f"AI analysis saved: {saved_vals.get('diagnosed_disease','')} · Risk: {saved_vals.get('risk_level','')}")
                                st.session_state['last_analysis'] = saved_vals
                            else:
                                st.warning("AI analysis completed but could not verify save to database.")
                            conn.close()

                            refreshed_patient = get_latest_patient_record(st.session_state.user_id)
                            report_data, report_name, report_mime = generate_patient_report(
                                {**(refreshed_patient or {}), **dict(user_info), 'record_date': (refreshed_patient or {}).get('created_at')},
                                doctor_prescriptions
                            )
                            st.download_button(
                                label='📄 Download Patient Health Report',
                                data=report_data,
                                file_name=report_name,
                                mime=report_mime,
                                use_container_width=True,
                                help='Download a complete patient report including AI diagnosis, drug recommendations, drug interaction warnings, risk analysis, and doctor prescription details.'
                            )
                    except Exception as e:
                        st.error(f"Analysis error: {e}")

    # ── Book Appointment ───────────────────────────────────────────────────────
    elif "Book Appointment" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Scheduling</div>
            <div class="hero-title">Book an <span>Appointment</span></div>
            <div class="hero-sub">Choose a verified specialist and pick a convenient time slot.</div>
        </div>
        """, unsafe_allow_html=True)

        conn = get_db()
        doctors = conn.execute("""
            SELECT u.id, u.name, d.specialization, d.hospital, d.experience_years, d.available_days, d.consultation_fee, d.rating
            FROM doctors d JOIN users u ON d.user_id=u.id
        """).fetchall()
        conn.close()

        if not doctors:
            st.info("No doctors registered yet. Check back later.")
        else:
            doc_options = {f"Dr. {d['name']} — {d['specialization']} ({d['hospital']})": d['id'] for d in doctors}
            sel_doc_label = st.selectbox("Select Doctor", list(doc_options.keys()))
            sel_doc_id    = doc_options[sel_doc_label]
            sel_doc       = next(d for d in doctors if d['id'] == sel_doc_id)

            st.markdown(f"""
            <div class="doc-card" style="margin:10px 0 20px;">
                <div class="doc-name">Dr. {sel_doc['name']}</div>
                <div class="doc-spec">{sel_doc['specialization']}</div>
                <div style="font-size:0.85rem;color:#64748b;margin-bottom:8px;">🏥 {sel_doc['hospital']} &nbsp;|&nbsp; {sel_doc['experience_years']} yrs experience</div>
                <div style="font-size:0.85rem;color:#94a3b8;margin-bottom:10px;">📅 Available: {sel_doc['available_days']}</div>
                <span class="badge badge-green">₹{sel_doc['consultation_fee']} / visit</span>
            </div>
            """, unsafe_allow_html=True)

            with st.form("book_form"):
                b1, b2 = st.columns(2)
                with b1:
                    appt_date = st.date_input("Appointment Date", min_value=date.today())
                with b2:
                    appt_time = st.time_input("Preferred Time")
                reason = st.text_area("Reason for Visit", placeholder="Describe your symptoms or concern…", height=100)
                if st.form_submit_button("Confirm Booking →", use_container_width=True):
                    conn = get_db()
                    conn.execute("INSERT INTO appointments(patient_id,doctor_id,appointment_date,appointment_time,reason) VALUES(?,?,?,?,?)",
                        (st.session_state.user_id, sel_doc_id, str(appt_date), str(appt_time), reason))
                    conn.commit()
                    conn.close()
                    st.success(f"✅ Appointment booked with Dr. {sel_doc['name']} on {appt_date} at {appt_time}.")

    # ── Find Doctors ───────────────────────────────────────────────────────────
    elif "Find Doctors" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Doctor Directory</div>
            <div class="hero-title">Find <span>Specialists</span></div>
            <div class="hero-sub">Browse our verified network of healthcare professionals.</div>
        </div>
        """, unsafe_allow_html=True)

        conn = get_db()
        doctors = conn.execute("""
            SELECT u.id, u.name, u.phone, d.specialization, d.hospital, d.experience_years,
                   d.available_days, d.consultation_fee, d.rating, d.total_reviews
            FROM doctors d JOIN users u ON d.user_id=u.id
        """).fetchall()
        conn.close()

        spec_filter = st.text_input("🔍 Search by specialization or name", placeholder="e.g. Cardiology, Dr. Sharma")
        filtered = [d for d in doctors if not spec_filter or
                    spec_filter.lower() in d['specialization'].lower() or
                    spec_filter.lower() in d['name'].lower()]

        if not filtered:
            st.info("No doctors found. Try a different search term.")
        for doc in filtered:
            st.markdown(f"""
            <div class="doc-card">
                <div style="display:flex;justify-content:space-between;align-items:flex-start;">
                    <div>
                        <div class="doc-name">Dr. {doc['name']}</div>
                        <div class="doc-spec">{doc['specialization']}</div>
                        <div style="font-size:0.85rem;color:#64748b;">🏥 {doc['hospital']} &nbsp;·&nbsp; {doc['experience_years']} yrs</div>
                        <div style="font-size:0.85rem;color:#64748b;margin-top:4px;">📅 {doc['available_days']} &nbsp;·&nbsp; 📞 {doc['phone'] or 'N/A'}</div>
                    </div>
                    <div style="text-align:right;">
                        <div style="font-family:'Syne',sans-serif;font-size:1.4rem;font-weight:800;color:#00d4ff;">₹{doc['consultation_fee']}</div>
                        <div style="font-size:0.75rem;color:#64748b;">per visit</div>
                        <div style="margin-top:6px;"><span class="badge badge-green">⭐ {doc['rating'] or 'New'}</span></div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

    # ── Medical History ─────────────────────────────────────────────────────────
    elif "History" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Records</div>
            <div class="hero-title">Medical <span>History</span></div>
            <div class="hero-sub">Your complete health timeline — diagnoses, vitals, and risk trends.</div>
        </div>
        """, unsafe_allow_html=True)

        conn = get_db()
        df = pd.read_sql_query("SELECT * FROM patients WHERE user_id=?", conn, params=(st.session_state.user_id,))
        user_info = conn.execute("SELECT name, age, gender, phone FROM users WHERE id=?", (st.session_state.user_id,)).fetchone()
        conn.close()

        if df.empty:
            st.info("No records yet. Enter medical data to get started.")
        else:
            m1, m2, m3, m4 = st.columns([1,1,1,1])
            m1.metric("Total Records", len(df))
            m2.metric("Unique Diagnoses", df['diagnosed_disease'].nunique())
            m3.metric("High Risk Records", len(df[df['risk_level'] == 'High']))

            latest_record = get_latest_patient_record(st.session_state.user_id)
            if latest_record:
                latest_prescriptions = []
                report_conn = get_db()
                latest_prescriptions = report_conn.execute(
                    "SELECT drug, dosage, frequency, duration, notes, date_prescribed FROM prescriptions WHERE patient_id=? ORDER BY date_prescribed DESC",
                    (latest_record['id'],)
                ).fetchall()
                report_conn.close()
                report_data, report_name, report_mime = generate_patient_report(
                    {**latest_record, **(dict(user_info) if user_info else {}), 'record_date': latest_record.get('created_at')},
                    latest_prescriptions
                )
                m4.download_button(
                    label='📄 Download Latest Health Report',
                    data=report_data,
                    file_name=report_name,
                    mime=report_mime,
                    use_container_width=True,
                    help='Download a complete health report from the latest saved medical record.'
                )
            else:
                m4.write("No saved history yet.")

            risk_counts = df['risk_level'].fillna('Unknown').astype(str).value_counts()
            risk_counts = risk_counts.reindex(['Low','Medium','High','Unknown'], fill_value=0)
            risk_chart = pd.DataFrame({
                'Risk Level': risk_counts.index,
                'Count': risk_counts.values
            }).set_index('Risk Level')
            st.bar_chart(risk_chart)
            st.markdown(
                f"<div style='margin-top:8px;color:#94a3b8;'>"
                f"The bar chart shows how many records are classified at each risk level. "
                f"Total records: <strong>{len(df)}</strong>. "
                f"Low: <strong>{risk_counts['Low']}</strong>, "
                f"Medium: <strong>{risk_counts['Medium']}</strong>, "
                f"High: <strong>{risk_counts['High']}</strong>, "
                f"Unknown: <strong>{risk_counts['Unknown']}</strong>.</div>",
                unsafe_allow_html=True
            )

            with st.expander("📋 Full Record Table"):
                display_df = df[[
                    'blood_pressure', 'blood_sugar', 'allergies', 'past_diseases', 'current_symptoms',
                    'previous_drugs', 'diagnosed_disease', 'recommended_drug', 'alternative_drug',
                    'side_effects', 'risk_level', 'created_at'
                ]].fillna('—')
                display_df = display_df.rename(columns={
                    'blood_pressure': 'Blood Pressure',
                    'blood_sugar': 'Blood Sugar',
                    'allergies': 'Allergies',
                    'past_diseases': 'Past Diseases',
                    'current_symptoms': 'Current Symptoms',
                    'previous_drugs': 'Previous Drugs',
                    'diagnosed_disease': 'Diagnosed Disease',
                    'recommended_drug': 'Recommended Drug',
                    'alternative_drug': 'Alternative Drug',
                    'side_effects': 'Side Effects',
                    'risk_level': 'Risk Level',
                    'created_at': 'Record Created'
                })
                st.dataframe(display_df, use_container_width=True)

            st.markdown('<div class="sec-title" style="margin:20px 0 12px;">🕐 Timeline</div>', unsafe_allow_html=True)
            for _, row in df.sort_values('id', ascending=False).iterrows():
                if row['diagnosed_disease']:
                    rb = {"High":"badge-red","Medium":"badge-yellow","Low":"badge-green"}.get(row['risk_level'],'badge-blue')
                    st.markdown(f"""
                    <div class="timeline-item">
                        <div class="timeline-dot"></div>
                        <div style="font-size:0.75rem;color:#64748b;">{row['created_at']}</div>
                        <div style="font-weight:600;color:#f1f5f9;margin:3px 0;">{row['diagnosed_disease']}</div>
                        <span class="badge {rb}">{row['risk_level']} Risk</span>
                        <span style="font-size:0.82rem;color:#64748b;margin-left:10px;">💊 {row['recommended_drug'] or '—'}</span>
                    </div>
                    """, unsafe_allow_html=True)

    # ── Prescriptions ───────────────────────────────────────────────────────────
    elif "Prescriptions" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Pharmacy</div>
            <div class="hero-title">My <span>Prescriptions</span></div>
            <div class="hero-sub">All drug recommendations and prescriptions in one place.</div>
        </div>
        """, unsafe_allow_html=True)

        conn = get_db()
        pres = conn.execute("""
            SELECT pr.drug, pr.dosage, pr.frequency, pr.duration, pr.date_prescribed, pr.notes
            FROM prescriptions pr
            JOIN patients pt ON pr.patient_id = pt.id
            WHERE pt.user_id=?
            ORDER BY pr.date_prescribed DESC
        """, (st.session_state.user_id,)).fetchall()
        conn.close()

        if not pres:
            st.info("No prescriptions yet. Run an AI analysis to get drug recommendations.")
        for p in pres:
            st.markdown(f"""
            <div class="mc-card mc-card-green">
                <div style="display:flex;justify-content:space-between;align-items:flex-start;">
                    <div>
                        <div style="font-family:'Syne',sans-serif;font-size:1.1rem;font-weight:700;color:#f1f5f9;">{p['drug']}</div>
                        <div style="font-size:0.85rem;color:#64748b;margin-top:4px;">📦 Dosage: {p['dosage'] or 'As prescribed'} &nbsp;·&nbsp; 🔁 {p['frequency'] or '—'} &nbsp;·&nbsp; ⏱ {p['duration'] or '—'}</div>
                        {f'<div style="font-size:0.82rem;color:#94a3b8;margin-top:6px;">📝 {p["notes"]}</div>' if p['notes'] else ''}
                    </div>
                    <div><span class="badge badge-blue">{p['date_prescribed'][:10] if p['date_prescribed'] else '—'}</span></div>
                </div>
            </div>
            """, unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# DOCTOR PAGES
# ══════════════════════════════════════════════════════════════════════════════
elif st.session_state.user_role == 'doctor':

    conn = get_db()
    doc_info = conn.execute("SELECT * FROM doctors WHERE user_id=?", (st.session_state.user_id,)).fetchone()
    conn.close()

    # ── Doctor Dashboard ───────────────────────────────────────────────────────
    if "Dashboard" in page:
        conn = get_db()
        total_appts  = conn.execute("SELECT COUNT(*) FROM appointments WHERE doctor_id=?", (st.session_state.user_id,)).fetchone()[0]
        pending      = conn.execute("SELECT COUNT(*) FROM appointments WHERE doctor_id=? AND status='pending'", (st.session_state.user_id,)).fetchone()[0]
        completed    = conn.execute("SELECT COUNT(*) FROM appointments WHERE doctor_id=? AND status='completed'", (st.session_state.user_id,)).fetchone()[0]
        today_appts  = conn.execute("SELECT a.appointment_time, u.name, a.reason, a.id, a.status FROM appointments a JOIN users u ON a.patient_id=u.id WHERE a.doctor_id=? AND a.appointment_date=date('now') ORDER BY a.appointment_time", (st.session_state.user_id,)).fetchall()
        conn.close()

        spec = doc_info['specialization'] if doc_info else '—'
        doctor_name = clean_doctor_display_name(st.session_state.user_name)
        st.markdown(f"""
        <div class="hero-wrap">
            <div class="hero-eyebrow">{spec} · Doctor Portal</div>
            <div class="hero-title">Welcome, Dr. <span>{doctor_name}</span></div>
            <div class="hero-sub">Manage your appointments, patients, and prescriptions from your personal dashboard.</div>
        </div>
        """, unsafe_allow_html=True)

        m1, m2, m3 = st.columns(3)
        m1.metric("Total Appointments", total_appts)
        m2.metric("Pending", pending)
        m3.metric("Completed", completed)

        st.markdown('<div class="sec-title" style="margin:24px 0 12px;">📅 Today\'s Schedule</div>', unsafe_allow_html=True)
        if not today_appts:
            st.info("No appointments scheduled for today.")
        for a in today_appts:
            status_badge = {"pending":"badge-yellow","confirmed":"badge-blue","completed":"badge-green","cancelled":"badge-red"}.get(a['status'],'badge-blue')
            st.markdown(f"""
            <div class="mc-card" style="padding:16px 20px;margin-bottom:10px;">
                <div style="display:flex;justify-content:space-between;align-items:center;">
                    <div>
                        <span style="font-size:0.8rem;color:#00d4ff;">🕐 {a['appointment_time']}</span>
                        <div style="font-weight:600;color:#f1f5f9;margin:3px 0;">{a['name']}</div>
                        <div style="font-size:0.85rem;color:#64748b;">{a['reason'] or '—'}</div>
                    </div>
                    <span class="badge {status_badge}">{a['status'].title()}</span>
                </div>
            </div>
            """, unsafe_allow_html=True)

    # ── Doctor Appointments ────────────────────────────────────────────────────
    elif "Appointments" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Schedule</div>
            <div class="hero-title">My <span>Appointments</span></div>
        </div>
        """, unsafe_allow_html=True)

        conn = get_db()
        appts = conn.execute("""
            SELECT a.id, a.appointment_date, a.appointment_time, a.reason, a.status, a.notes, u.name, u.age, u.gender
            FROM appointments a JOIN users u ON a.patient_id=u.id
            WHERE a.doctor_id=? ORDER BY a.appointment_date DESC
        """, (st.session_state.user_id,)).fetchall()
        conn.close()

        status_filter = st.selectbox("Filter by status", ["All", "pending", "confirmed", "completed", "cancelled"])
        filtered_appts = [a for a in appts if status_filter == "All" or a['status'] == status_filter]

        for a in filtered_appts:
            sb = {"pending":"badge-yellow","confirmed":"badge-blue","completed":"badge-green","cancelled":"badge-red"}.get(a['status'],'badge-blue')
            with st.expander(f"📋 {a['name']} · {a['appointment_date']} {a['appointment_time']}"):
                c1, c2 = st.columns([2,1])
                with c1:
                    st.markdown(f"**Patient:** {a['name']} ({a['gender']}, {a['age']} yrs)")
                    st.markdown(f"**Reason:** {a['reason'] or '—'}")
                    st.markdown(f"**Notes:** {a['notes'] or '—'}")
                with c2:
                    st.markdown(f'<span class="badge {sb}">{a["status"].title()}</span>', unsafe_allow_html=True)
                    new_status = st.selectbox("Update Status", ["pending","confirmed","completed","cancelled"], key=f"st_{a['id']}", index=["pending","confirmed","completed","cancelled"].index(a['status']))
                    notes_input = st.text_input("Add Notes", key=f"nt_{a['id']}", value=a['notes'] or '')
                    if st.button("Update", key=f"upd_{a['id']}"):
                        conn = get_db()
                        conn.execute("UPDATE appointments SET status=?, notes=? WHERE id=?", (new_status, notes_input, a['id']))
                        conn.commit()
                        conn.close()
                        st.success("Updated!"); st.rerun()

    # ── My Patients ────────────────────────────────────────────────────────────
    elif "Patients" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Patient Management</div>
            <div class="hero-title">My <span>Patients</span></div>
            <div class="hero-sub">View the latest clinical summary for each patient who has booked with you.</div>
        </div>
        """, unsafe_allow_html=True)

        conn = get_db()
        patients = conn.execute("""
            SELECT DISTINCT u.id, u.name, u.age, u.gender, u.phone,
                   p.diagnosed_disease, p.risk_level, p.blood_pressure, p.blood_sugar,
                   p.recommended_drug, p.alternative_drug, p.side_effects, p.created_at AS record_date
            FROM appointments a
            JOIN users u ON a.patient_id=u.id
            LEFT JOIN patients p ON p.id = (
                SELECT MAX(id)
                FROM patients p2
                WHERE p2.user_id = u.id
            )
            WHERE a.doctor_id=?
            ORDER BY u.name
        """, (st.session_state.user_id,)).fetchall()
        conn.close()

        if not patients:
            st.info("No patients have booked with you yet.")
        for p in patients:
            rb = {"High":"badge-red","Medium":"badge-yellow","Low":"badge-green"}.get(p['risk_level'],'badge-blue')
            risk_html = f'<div style="margin-top:8px;"><span class="badge {rb}">Risk: {p["risk_level"]}</span></div>' if p['risk_level'] else ''
            diagnosis_html = f'<div style="font-size:0.85rem;color:#f1f5f9;font-weight:600;">{p["diagnosed_disease"]}</div>' if p['diagnosed_disease'] else '<div style="color:#64748b;font-size:0.85rem;">No diagnosis</div>'
            recommended_html = f'<div style="margin-top:8px;color:#cbd5e1;font-size:0.85rem;">Recommended: {p["recommended_drug"]}</div>' if p['recommended_drug'] else ''
            alternative_html = f'<div style="margin-top:4px;color:#cbd5e1;font-size:0.85rem;">Alternative: {p["alternative_drug"]}</div>' if p['alternative_drug'] else ''
            side_effects_html = f'<div style="margin-top:4px;color:#cbd5e1;font-size:0.85rem;">Side effects: {p["side_effects"]}</div>' if p['side_effects'] else ''
            card_html = textwrap.dedent(f"""
            <div class="mc-card" style="margin-bottom:12px;padding:18px;">
                <div style="display:flex;justify-content:space-between;flex-wrap:wrap;gap:16px;">
                    <div style="min-width:260px;max-width:calc(100% - 220px);">
                        <div style="font-family:'Syne',sans-serif;font-weight:700;font-size:1.05rem;color:#f1f5f9;">{p['name']}</div>
                        <div style="font-size:0.85rem;color:#64748b;">{p['gender']}, {p['age']} yrs · 📞 {p['phone'] or 'N/A'}</div>
                        <div style="font-size:0.85rem;color:#94a3b8;margin-top:4px;">BP: {p['blood_pressure'] or '—'} · Sugar: {p['blood_sugar'] or '—'} mg/dL</div>
                        {risk_html}
                        {recommended_html}
                        {alternative_html}
                        {side_effects_html}
                    </div>
                    <div style="min-width:180px;max-width:240px;text-align:right;flex:1;">
                        {diagnosis_html}
                        <div style="font-size:0.75rem;color:#94a3b8;margin-top:8px;">Latest record: {p['record_date'][:10] if p['record_date'] else '—'}</div>
                    </div>
                </div>
            </div>
            """)
            st.markdown(card_html, unsafe_allow_html=True)
            record_id = get_latest_patient_record_id(p['id'])
            prescriptions = []
            patient_record = None
            if record_id:
                conn = get_db()
                prescriptions = conn.execute(
                    "SELECT drug, dosage, frequency, duration, notes, date_prescribed FROM prescriptions WHERE patient_id=? ORDER BY date_prescribed DESC",
                    (record_id,)
                ).fetchall()
                patient_row = conn.execute("SELECT * FROM patients WHERE id=?", (record_id,)).fetchone()
                user_row = conn.execute("SELECT name, age, gender, phone FROM users WHERE id=?", (p['id'],)).fetchone()
                conn.close()
                if patient_row:
                    patient_record = dict(patient_row)
                    if user_row:
                        patient_record.update(dict(user_row))
            if patient_record:
                report_bytes, report_name, report_mime = generate_patient_report(patient_record, prescriptions)
            else:
                # Fallback: generate report from user info if no patient record exists
                fallback = {**dict(p), 'record_date': 'N/A'}
                report_bytes, report_name, report_mime = generate_patient_report(fallback, prescriptions)
            st.download_button(
                label='Download Patient Report',
                data=report_bytes,
                file_name=report_name,
                mime=report_mime,
                key=f'download_report_{p["id"]}',
                help='Generate a document summary of patient diagnosis, AI recommendations, and doctor prescriptions.'
            )

    # ── Write Prescription ─────────────────────────────────────────────────────
    elif "Prescription" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Clinical Tools</div>
            <div class="hero-title">Write a <span>Prescription</span></div>
        </div>
        """, unsafe_allow_html=True)

        conn = get_db()
        my_patients = conn.execute("""
            SELECT DISTINCT u.id, u.name
            FROM appointments a JOIN users u ON a.patient_id=u.id
            WHERE a.doctor_id=?
            ORDER BY u.name
        """, (st.session_state.user_id,)).fetchall()
        conn.close()

        if not my_patients:
            st.info("No patients found. Patients need to book appointments with you first.")
        else:
            pat_map = {}
            for p in my_patients:
                uid_p = p['id']
                pid_p = get_latest_patient_record_id(uid_p)
                pat_map[f"{p['name']}"] = (uid_p, pid_p)
            sel_pat = st.selectbox("Select Patient", list(pat_map.keys()))
            uid, pid = pat_map[sel_pat]

            with st.form("rx_form"):
                r1, r2 = st.columns(2)
                with r1:
                    drug     = st.text_input("Drug Name")
                    dosage   = st.text_input("Dosage", placeholder="e.g. 500mg")
                    frequency= st.selectbox("Frequency", ["Once daily","Twice daily","Three times daily","As needed","Before meals","After meals"])
                with r2:
                    duration = st.text_input("Duration", placeholder="e.g. 7 days")
                    notes    = st.text_area("Doctor's Notes", height=90)

                if st.form_submit_button("Issue Prescription →", use_container_width=True):
                    if not drug:
                        st.error("Please enter a drug name.")
                    else:
                        # Check interactions with patient's existing drugs
                        conn = get_db()
                        pat_rec = conn.execute("SELECT previous_drugs FROM patients WHERE id=?", (pid,)).fetchone()
                        conn.close()

                        if pid is None:
                            st.error("Selected patient has no medical record. Ask the patient to add medical data before issuing a prescription.")
                        else:
                            if pat_rec and pat_rec['previous_drugs']:
                                prev = [d.strip() for d in pat_rec['previous_drugs'].split(',') if d.strip()]
                                warn_list = []
                                for pd_ in prev:
                                    inter, sev = check_interaction(drug, pd_)
                                    if inter:
                                        warn_list.append(f"{drug} + {pd_}: {inter} ({sev})")
                                if warn_list:
                                    st.warning("⚠️ Drug Interaction Warning:\n" + "\n".join(warn_list))

                            conn = get_db()
                            conn.execute("INSERT INTO prescriptions(patient_id,doctor_id,drug,dosage,frequency,duration,notes,date_prescribed) VALUES(?,?,?,?,?,?,?,?)",
                                (pid, st.session_state.user_id, drug, dosage, frequency, duration, notes, datetime.now().strftime("%Y-%m-%d")))
                            conn.commit()
                            conn.close()
                            st.success(f"✅ Prescription for {drug} issued to {sel_pat}.")

    # ── Analytics ──────────────────────────────────────────────────────────────
    elif "Analytics" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Insights</div>
            <div class="hero-title">Practice <span>Analytics</span></div>
        </div>
        """, unsafe_allow_html=True)

        conn = get_db()
        appt_df = pd.read_sql_query("SELECT * FROM appointments WHERE doctor_id=?", conn, params=(st.session_state.user_id,))
        conn.close()

        if appt_df.empty:
            st.info("No data yet to show analytics.")
        else:
            m1, m2, m3 = st.columns(3)
            m1.metric("Total Appointments", len(appt_df))
            m2.metric("Completion Rate", f"{len(appt_df[appt_df['status']=='completed'])/len(appt_df)*100:.0f}%")
            m3.metric("Pending", len(appt_df[appt_df['status']=='pending']))

            st.markdown('<div class="sec-title" style="margin:20px 0 10px;">Appointments by Status</div>', unsafe_allow_html=True)
            st.bar_chart(appt_df['status'].value_counts())

    # ── Doctor Profile ─────────────────────────────────────────────────────────
    elif "Profile" in page:
        st.markdown("""
        <div class="hero-wrap">
            <div class="hero-eyebrow">Settings</div>
            <div class="hero-title">Doctor <span>Profile</span></div>
        </div>
        """, unsafe_allow_html=True)

        if doc_info:
            conn = get_db()
            u = conn.execute("SELECT * FROM users WHERE id=?", (st.session_state.user_id,)).fetchone()
            conn.close()
            st.markdown(f"""
            <div class="mc-card mc-card-accent">
                <div style="font-family:'Syne',sans-serif;font-size:1.5rem;font-weight:800;color:#f1f5f9;margin-bottom:4px;">Dr. {u['name']}</div>
                <div style="color:#00d4ff;margin-bottom:16px;">{doc_info['specialization']}</div>
                <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;">
                    <div><span style="font-size:0.75rem;color:#64748b;">HOSPITAL</span><br><span style="color:#e2e8f0;">{doc_info['hospital'] or '—'}</span></div>
                    <div><span style="font-size:0.75rem;color:#64748b;">LICENSE</span><br><span style="color:#e2e8f0;">{doc_info['license_no'] or '—'}</span></div>
                    <div><span style="font-size:0.75rem;color:#64748b;">EXPERIENCE</span><br><span style="color:#e2e8f0;">{doc_info['experience_years']} years</span></div>
                    <div><span style="font-size:0.75rem;color:#64748b;">FEE</span><br><span style="color:#e2e8f0;">₹{doc_info['consultation_fee']}</span></div>
                    <div><span style="font-size:0.75rem;color:#64748b;">AVAILABLE</span><br><span style="color:#e2e8f0;">{doc_info['available_days'] or '—'}</span></div>
                    <div><span style="font-size:0.75rem;color:#64748b;">EMAIL</span><br><span style="color:#e2e8f0;">{u['email']}</span></div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            st.markdown('<div class="sec-title" style="margin:20px 0 12px;">Update Profile</div>', unsafe_allow_html=True)
            with st.form("profile_form"):
                p1, p2 = st.columns(2)
                with p1:
                    new_hosp = st.text_input("Hospital", value=doc_info['hospital'] or '')
                    new_days = st.text_input("Available Days", value=doc_info['available_days'] or '')
                with p2:
                    new_fee  = st.number_input("Consultation Fee (₹)", value=float(doc_info['consultation_fee'] or 500))
                    new_exp  = st.number_input("Experience (Years)", value=int(doc_info['experience_years'] or 0))
                if st.form_submit_button("Save Changes →", use_container_width=True):
                    conn = get_db()
                    conn.execute("UPDATE doctors SET hospital=?,available_days=?,consultation_fee=?,experience_years=? WHERE user_id=?",
                        (new_hosp, new_days, new_fee, new_exp, st.session_state.user_id))
                    conn.commit()
                    conn.close()
                    st.success("Profile updated!"); st.rerun()
        else:
            st.error("Doctor profile not found. Please contact support.")